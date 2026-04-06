#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate bilingual (Chinese + English) Disclaimer statement as Word document.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_disclaimer_docx():
    """Create bilingual disclaimer Word document."""

    doc = Document()

    # Title - Chinese
    title_cn = doc.add_heading('免责声明', 0)
    title_run_cn = title_cn.runs[0]
    title_run_cn.font.size = Pt(24)
    title_run_cn.font.bold = True
    title_run_cn.font.color.rgb = RGBColor(0, 100, 160)

    # Title - English
    title_en = doc.add_heading('DISCLAIMER', 0)
    title_run_en = title_en.runs[0]
    title_run_en.font.size = Pt(24)
    title_run_en.font.bold = True
    title_run_en.font.color.rgb = RGBColor(0, 100, 160)

    doc.add_paragraph()

    # Separator line
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('=' * 60)
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Chinese content
    cn_section_title = doc.add_heading('中文声明', level=1)
    cn_run = cn_section_title.runs[0]
    cn_run.font.size = Pt(18)
    cn_run.font.bold = True
    cn_run.font.color.rgb = RGBColor(0, 80, 120)

    disclaimer_cn_text = (
        "地球村民监察互助联盟榜单免责声明\n"
        "\n"
        "本榜单旨在社会公益监督，实时展示公众投票结果，不构成事实认定，不针对任何企事业机构和个人，无法律效力，亦不承担法律责任，用户需自行独立思考，仔细甄辨。\n"
        "\n"
        "对于未成年人或缺乏甄辨能力的个人，建议在监护人指导下进行思考决策。\n"
        "\n"
        "本榜采用防刷票算法机制，并采取加密机制保护用户隐私。生态系统由公众共同搭建和维护，系统的健康和可持续发展靠民智觉悟和努力奋斗！\n"
    )

    para = doc.add_paragraph()
    para.add_run(disclaimer_cn_text).font.size = Pt(12)
    para.paragraph_format.space_after = Pt(20)

    # English content
    en_section_title = doc.add_heading('English Statement', level=1)
    en_run = en_section_title.runs[0]
    en_run.font.size = Pt(18)
    en_run.font.bold = True
    en_run.font.color.rgb = RGBColor(0, 80, 120)

    disclaimer_en_text = (
        "Earth Citizens Supervision Alliance - Disclaimer\n"
        "\n"
        "This leaderboard is designed for public welfare supervision and displays real-time voting results from the public. It does not constitute a factual determination, is not targeted at any enterprise or individual, has no legal effect, and does not assume any legal liability. Users must think independently and carefully verify information.\n"
        "\n"
        "For minors or individuals lacking verification capabilities, it is recommended to make decisions under the guidance of guardians.\n"
        "\n"
        "This leaderboard employs anti-fraud voting algorithms and encryption mechanisms to protect user privacy. The ecosystem is built and maintained collectively by the public, and its healthy and sustainable development relies on public wisdom, awareness, and hard work!\n"
    )

    para = doc.add_paragraph()
    para.add_run(disclaimer_en_text).font.size = Pt(12)
    para.paragraph_format.space_after = Pt(20)

    # Footer section
    doc.add_paragraph()
    footer_line = doc.add_paragraph()
    run = footer_line.add_run('-' * 60)
    run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_para = doc.add_paragraph()
    footer_text = "文档生成日期：2026-04-06 | Generated: 2026-04-06\n项目地址：地球村民监察互助联盟"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(footer_text)
    run.font.size = Pt(10)

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), '免责声明.docx')
    doc.save(output_path)
    print(f'[OK] 双语免责声明文档已保存：{output_path}')


def main():
    """Main entry point."""
    print('=' * 60)
    print('Bilingual Disclaimer Generator - Earth Citizens Supervision Alliance')
    print('=' * 60)

    create_disclaimer_docx()

    print('\n' + '=' * 60)
    print('Disclaimer successfully generated!')
    print('=' * 60)


if __name__ == '__main__':
    main()
