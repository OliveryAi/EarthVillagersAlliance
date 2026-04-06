#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档生成脚本 - 直接创建 Word 和 PDF 格式文档
无需依赖 markdown 库
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


def create_word_docx():
    """创建 Word 文档 (开发过程记录)"""

    doc = Document()
    doc.add_heading('地球村民监察互助联盟 - 项目完整开发过程记录', 0)

    # Section 1
    doc.add_heading('第一部分：初始架构设计与规划', level=1)

    p = doc.add_paragraph()
    p.add_run('项目名称').bold = True
    p.add_run(': 地球村民监察互助联盟\n')
    p.add_run('核心目标').bold = True
    p.add_run(': 建立网友投票监督平台，形成就业歧视/剥削行为榜单，推动社会透明化\n')

    # Table: Tech Stack
    doc.add_heading('1.1 技术栈选择', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['组件', '技术方案', '选择理由']
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True

    data = [
        ['Web 框架', 'Django + DRF', '内置用户认证、后台管理、ORM'],
        ['数据库', 'SQLite', '轻量级无需部署，便于本地测试'],
        ['GUI 客户端', 'PyQt6', '功能丰富的跨平台桌面应用框架'],
        ['安全机制', 'TokenAuthentication + 设备指纹', '保障投票真实性']
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text

    # Section 2
    doc.add_heading('第二部分：数据库模型开发过程', level=1)
    doc.add_heading('2.1 初始模型定义', level=2)

    code_block = doc.add_paragraph()
    runner = code_block.add_run('''class Evidence(models.Model):
    """证据材料表"""
    vote = models.ForeignKey(Vote, related_name='evidences', on_delete=models.CASCADE)
    file = models.CharField(max_length=100, verbose_name='文件路径')  # 字段命名不够规范
    description = models.TextField(verbose_name='描述')
''')
    runner.font.name = 'Courier New'

    doc.add_heading('2.2 模型修复过程', level=2)

    p = doc.add_paragraph()
    p.add_run('发现').bold = True
    p.add_run('：Evidence 表使用了 file、description、is_verified 字段，与预期设计不符。\n')

    p = doc.add_paragraph()
    p.add_run('最终模型定义 (v1.1):').italic = True

    code_block = doc.add_paragraph()
    runner = code_block.add_run('''class Evidence(models.Model):
    """证据材料表 - 用户上传的截图/文档"""
    vote = models.ForeignKey(Vote, related_name='evidences', on_delete=models.CASCADE)
    file_path = models.CharField(max_length=500, verbose_name='文件路径')
    file_type = models.CharField(max_length=20, choices=[('image', '图片'), ('document', '文档')])

    class Meta:
        db_table = 'evidence'
''')
    runner.font.name = 'Courier New'

    # Section 3
    doc.add_heading('第三部分：用户认证系统开发过程', level=1)
    doc.add_heading('3.1 CustomUser 与 Native User 分离架构', level=2)

    p = doc.add_paragraph()
    p.add_run('设计决策').bold = True
    p.add_run('\n')
    p.add_run('• ').bold = True
    p.add_run('CustomUser: 存储加密手机号、验证码等敏感信息\n')
    p.add_run('• ').bold = True
    p.add_run('Native User (django.contrib.auth.User): DRF TokenAuthentication 必须关联到此模型')

    doc.add_heading('3.2 Username Bridge Pattern', level=2)

    p = doc.add_paragraph()
    p.add_run('关键代码段 (views.py Line 50-60)').italic = True

    code_block = doc.add_paragraph()
    runner = code_block.add_run('''from django.contrib.auth import get_user_model
NativeUser = get_user_model()

if isinstance(request.user, NativeUser):
    custom_user = CustomUser.objects.filter(username=request.user.username).first()
    if not custom_user:
        return Response({'error': '用户不存在'})
else:
    custom_user = request.user
''')
    runner.font.name = 'Courier New'

    doc.add_heading('第四部分：投票核心逻辑开发过程', level=1)
    doc.add_heading('4.1 防刷票机制实现', level=2)

    code_block = doc.add_paragraph()
    runner = code_block.add_run('''# IP/设备频率限制 (SubmitVoteView.py Line 66-84)
ip_address = request.META.get('REMOTE_ADDR') or request.META.get('HTTP_X_FORWARDED_FOR', '')
fingerprint = get_device_fingerprint(request)

daily_votes = Vote.objects.filter(
    voter_id=custom_user.id,
    category=category,
    ip_address__startswith=ip_address.split(',')[0],
    created_at__gte=today_midnight
).count()

if daily_votes >= RATE_LIMIT_PER_DAY:
    return Response({'error': '今日投票次数已达上限'}, status=429)
''')
    runner.font.name = 'Courier New'

    # Section 5 - Evidence Upload Process
    doc.add_heading('第五部分：证据上传功能开发过程', level=1)
    doc.add_heading('5.1 问题发现与解决', level=2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    headers = ['阶段', '结果']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    rows_data = [
        ['上传 PDF 证据 (第一次)', 'OperationalError: table evidence has no column named file_path'],
        ['原因分析', 'Django ORM 进程缓存旧模型定义'],
        ['解决方案', 'taskkill python.exe + runserver 重启服务器'],
        ['上传证据 (第二次)', 'NOT NULL constraint failed: evidence.file'],
    ]

    for i, row_data in enumerate(rows_data):
        row = table.rows[i + 1].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text

    doc.add_heading('第六部分：完整端到端测试总结', level=1)

    test_results = [
        ['序号', '场景', '预期结果', '实际结果'],
        ['1', '企业搜索（华为）', '返回 Huawei Technologies', '✅ 通过'],
        ['2', '提交年龄歧视投票', 'Vote ID=5, 成功创建', '✅ 通过'],
        ['3', '同一用户重复投票', '400 Bad Request', '✅ 通过'],
        ['4', 'IP 超限投票（第 4 次）', '429 Too Many Requests', '✅ 通过'],
        ['5', '上传 PDF 证据', '成功，返回 file_url', '✅ 通过'],
        ['6', '未登录用户访问', '401 Unauthorized', '✅ 通过'],
    ]

    table = doc.add_table(rows=len(test_results), cols=4)
    table.style = 'Table Grid'

    for i, row_data in enumerate(test_results):
        row = table.rows[i].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text

    # Statistics section
    doc.add_heading('统计数据', level=2)
    p = doc.add_paragraph()
    p.add_run('总投票数').bold = True
    p.add_run(': 4\n')
    p.add_run('证据上传数').bold = True
    p.add_run(': 2\n')
    p.add_run('年龄歧视榜票数').bold = True
    p.add_run(': 3\n')
    p.add_run('PUA 专制榜票数').bold = True
    p.add_run(': 1\n')

    doc.save(os.path.join(os.path.dirname(__file__), '项目开发过程记录.docx'))
    print('[OK] Word 文档已生成：项目开发过程记录.docx')


def create_pdf():
    """创建 PDF 文档"""

    from reportlab.lib import colors

    pdf_file = os.path.join(os.path.dirname(__file__), '项目开发过程记录.pdf')
    doc_build = SimpleDocTemplate(pdf_file, pagesize=A4)

    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=20, spacing=20)
    h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=16, spacing=14)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=9, leading=8)

    elements = []

    # Title
    elements.append(Paragraph('地球村民监察互助联盟 - 项目完整开发过程记录', h1_style))
    elements.append(Spacer(1, 20))

    # Section headers and content
    sections = [
        ('第一部分：初始架构设计与规划', '项目名称：地球村民监察互助联盟\n核心目标：建立网友投票监督平台，形成就业歧视/剥削行为榜单，推动社会透明化'),
        ('第二部分：数据库模型开发过程', '证据表字段修复 - 从 file/description/is_verified 迁移到 file_path/file_type'),
        ('第三部分：用户认证系统开发过程', 'CustomUser/NativeUser 分离架构，Username Bridge Pattern 实现'),
        ('第四部分：投票核心逻辑开发过程', '防刷票机制：IP+ 设备指纹双重验证，单日最多 3 票'),
        ('第五部分：证据上传功能开发过程', 'SQL ALTER TABLE 迁移旧数据，添加新字段并删除冗余列'),
        ('第六部分：完整端到端测试总结', '全部测试项通过 (6/6)，系统稳定运行'),
    ]

    for title, content in sections:
        elements.append(Paragraph(title, h2_style))
        elements.append(Paragraph(content, body_style))
        elements.append(Spacer(1, 10))

    # Table for test results
    from reportlab.platypus import Table as RLTable, TableStyle as RLTableStyle

    test_data = [
        ['序号', '场景', '预期结果', '实际结果'],
        ['1', '企业搜索（华为）', '返回 Huawei Technologies', '✅ 通过'],
        ['2', '提交投票', 'Vote ID=5, 成功创建', '✅ 通过'],
        ['3', '重复投票', '400 Bad Request', '✅ 通过'],
        ['4', '证据上传 PDF', '返回 file_url', '✅ 通过'],
    ]

    test_table = RLTable(test_data, colWidths=[50, 120, 100, 80])
    style = RLTableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ])
    test_table.setStyle(style)

    elements.append(test_table)

    doc_build.build(elements)
    print('[OK] PDF 文档已生成：项目开发过程记录.pdf')


def create_user_manual():
    """创建用户手册 (Word + PDF)"""

    # Word version
    doc = Document()
    doc.add_heading('地球村民监察互助联盟 - 用户操作手册', 0)

    sections = [
        ('目录', '1. 系统简介\n2. 快速开始\n3. 功能模块详解\n4. 常见问题解答'),
        ('1. 系统简介', '''四个观察榜：
• 35 岁以上就业歧视观察榜 (age_discrimination)
• 性别歧视观察榜 (gender_discrimination)
• 职场 PUA 专制观察榜 (pua_despotism)
• 过劳剥削观察榜 (overwork_exploitation)'''),
        ('2. 快速开始', '''步骤 1: 注册账号 - 手机号 +6 位验证码
步骤 2: 登录系统 - 获取 Token
步骤 3: 提交投票 - 选择榜单/企业，上传证据
步骤 4: 查看榜单排行 - Top N 展示'''),
        ('3. API 使用', '''POST /api/auth/register/ - 注册
POST /api/auth/login/ - 登录获取 Token
POST /api/vote/submit/ - 提交投票
GET /api/vote/ranking/<category>/ - 查看榜单
POST /api/vote/evidence/upload/ - 上传证据'''),
    ]

    for title, content in sections:
        doc.add_heading(title, level=1)
        for line in content.split('\n'):
            if line.startswith('• '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line.strip())

    doc.save(os.path.join(os.path.dirname(__file__), '使用说明文档.docx'))
    print('[OK] 用户手册 Word 已生成：使用说明文档.docx')

    # PDF version (simplified)
    pdf_file = os.path.join(os.path.dirname(__file__), '使用说明文档.pdf')
    doc_build = SimpleDocTemplate(pdf_file, pagesize=A4)
    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('ManualH1', parent=styles['Heading1'], fontSize=20)

    elements = [Paragraph('地球村民监察互助联盟 - 用户操作手册', h1_style)]

    for title, content in sections:
        elements.append(Spacer(1, 15))
        elements.append(Paragraph(title, styles['Heading2']))
        for line in content.split('\n'):
            if line.strip():
                elements.append(Paragraph(line.strip(), styles['Normal']))

    doc_build.build(elements)
    print('[OK] 用户手册 PDF 已生成：使用说明文档.pdf')


def main():
    """主函数"""
    print('='*60)
    print('文档生成程序 - 地球村民监察互助联盟')
    print('='*60)

    create_word_docx()
    create_pdf()
    create_user_manual()

    print('\n' + '='*60)
    print('所有文档已成功保存至项目根目录！')
    print('='*60)


if __name__ == '__main__':
    main()
