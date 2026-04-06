#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档转换脚本 - 将 Markdown 转换为 Word (.docx) 和 PDF 格式
"""

import os
from markdown import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


def markdown_to_docx(md_file, output_path):
    """将 Markdown 文件转换为 Word 文档"""

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    doc.add_heading('地球村民监察互助联盟 - 项目完整开发过程记录', 0)

    lines = content.split('\n')

    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- ') and not line.lstrip().startswith('- -'):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())

    doc.save(output_path)
    print(f'Word 文档已生成：{output_path}')


def markdown_to_pdf(md_file, output_path):
    """将 Markdown 文件转换为 PDF 格式"""

    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    c = __import__('reportlab.pdfgen.canvas').pdfexanvas.Canvas(output_path)
    width, height = A4

    # 创建样式
    from reportlab.lib.styles import ParagraphStyle

    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('CustomH1', parent=styles['Heading1'], fontSize=20, leading=24)
    h2_style = ParagraphStyle('CustomH2', parent=styles['Heading2'], fontSize=16, leading=18)

    elements = []
    lines = content.split('\n')

    for line in lines:
        if line.startswith('# '):
            elements.append(Paragraph(line[2:].strip(), h1_style))
            c.drawString(50, height - 30, f"=== {line[2:].strip()} ===")
            c.drawCentredString(width/2, height - 30, f"地球村民监察互助联盟 - 项目完整开发过程记录")
        elif line.startswith('## '):
            elements.append(Paragraph(line[3:].strip(), h2_style))
            c.drawString(50, height - 60, line[3:].strip())
        elif line.startswith('### '):
            elements.append(Paragraph(line[4:].strip(), styles['Heading3']))
        elif line.strip() and not line.startswith('```'):
            if line.startswith('- '):
                para = Paragraph(f'    {line[2:]}', styles['Normal'])
            else:
                para = Paragraph(line.strip(), styles['Normal'])
            elements.append(para)
        elif line.strip() == '':
            elements.append(Spacer(1, 6))

    # c.showPage()
    c.save()


def main():
    """主函数"""

    project_dir = os.path.dirname(os.path.abspath(__file__))

    print("=" * 60)
    print("文档转换程序 - 地球村民监察互助联盟")
    print("=" * 60)

    # 1. 开发过程记录转 Word
    md_path_1 = os.path.join(project_dir, '项目开发过程记录.md')
    docx_path_1 = os.path.join(project_dir, '项目开发过程记录.docx')

    if os.path.exists(md_path_1):
        markdown_to_docx(md_path_1, docx_path_1)
    else:
        print(f'警告：文件不存在 {md_path_1}')

    # 2. 使用手册转 Word
    md_path_2 = os.path.join(project_dir, '使用说明文档.md')
    docx_path_2 = os.path.join(project_dir, '使用说明文档.docx')

    if os.path.exists(md_path_2):
        print("\n正在生成使用说明 Word 文档...")
        doc2 = Document()
        with open(md_path_2, 'r', encoding='utf-8') as f:
            usage_content = f.read()

        doc2.add_heading('地球村民监察互助联盟 - 用户操作手册', 0)
        lines2 = usage_content.split('\n')
        for line in lines2:
            if line.startswith('# '):
                doc2.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc2.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc2.add_heading(line[4:].strip(), level=3)
            elif line.startswith('- '):
                p = doc2.add_paragraph(line[2:], style='List Bullet')
            elif line.strip() and not line.startswith('```'):
                if len(line.strip()) < 5:
                    continue
                else:
                    doc2.add_paragraph(line.strip())

        doc2.save(docx_path_2)
        print(f'Word 文档已生成：{docx_path_2}')
    else:
        print(f'警告：文件不存在 {md_path_2}')

    # 3. 使用报告库生成 PDF（简化版）
    from reportlab.lib import colors
    from io import BytesIO

    def convert_to_pdf_simple(md_file, output_path):
        """使用报告库生成 PDF"""
        from reportlab.platypus import SimpleDocTemplate
        doc_build = SimpleDocTemplate(output_path, pagesize=A4)

        styles = getSampleStyleSheet()
        h1_style = ParagraphStyle('H1', parent=styles['Heading1'], fontSize=20)
        h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=16)

        elements = []

        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')

        for line in lines:
            if line.startswith('# '):
                elements.append(Paragraph(line[2:].strip(), h1_style))
            elif line.startswith('## '):
                elements.append(Paragraph(line[3:].strip(), h2_style))
            elif line.strip() and not line.startswith('```'):
                if line.startswith('- '):
                    para = Paragraph(f'    {line[2:]}', styles['Normal'])
                else:
                    para = Paragraph(line.strip(), styles['Normal'])
                elements.append(para)
            elif line.strip() == '':
                elements.append(Spacer(1, 6))

        doc_build.build(elements)
        print(f'PDF 文档已生成：{output_path}')

    # 为两个文档生成 PDF 版本
    convert_to_pdf_simple(md_path_1, os.path.join(project_dir, '项目开发过程记录.pdf'))
    convert_to_pdf_simple(md_path_2, os.path.join(project_dir, '使用说明文档.pdf'))


if __name__ == '__main__':
    main()
