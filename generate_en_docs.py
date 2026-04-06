#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate English User Manual in Word and PDF formats.
"""

import os
from docx import Document


def create_en_word_docx():
    """Create English User Manual as Word document."""

    doc = Document()
    doc.add_heading('Earth Citizens Supervision Alliance - User Guide (English)', 0)

    # Table of Contents
    doc.add_heading('Table of Contents', level=1)

    toc_sections = [
        'System Overview',
        'Quick Start',
        'Feature Modules',
        'API Documentation Details',
        'FAQ & Troubleshooting',
        'Appendix A: Test Script Examples',
    ]

    for section in toc_sections:
        doc.add_paragraph(section, style='List Bullet')

    doc.add_page_break()

    # Section 1: System Overview
    doc.add_heading('1. System Overview', level=1)

    p = doc.add_paragraph()
    p.add_run('Project Introduction').bold = True
    p.add_run('\nEarth Citizens Supervision Alliance is a netizen-driven platform designed to monitor and document employment discrimination and labor exploitation practices through anonymous voting with evidence submission.\n')

    doc.add_heading('Four Observation Categories', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Category', 'Code Name', 'Description']
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True

    category_data = [
        ['35+ Age Discrimination', 'age_discrimination', 'Records enterprises with age-based hiring discrimination'],
        ['Gender Discrimination', 'gender_discrimination', 'Records enterprises with gender-based employment discrimination'],
        ['Workplace PUA Despotism', 'pua_despotism', 'Psychological manipulation and authoritarian management observation'],
        ['Overwork Exploitation', 'overwork_exploitation', 'Forced overtime and unpaid labor practices observation'],
    ]

    for i, row_data in enumerate(category_data):
        row = table.rows[i + 1].cells
        for j, cell_text in enumerate(row_data):
            row[j].text = cell_text

    doc.add_heading('Core Principles', level=2)
    principles = [
        'Anonymous Voting: All voter information is strictly confidential; only statistical results are published',
        'Evidence Priority: Users are encouraged to upload screenshots/documents as supporting evidence',
        'Anti-Fraud Mechanisms: Limited to 3 votes per user/day, one vote per company/category combination',
        'Real & Reliable: Device fingerprinting and IP rate limiting ensure data authenticity',
    ]

    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')

    doc.add_page_break()

    # Section 2: Quick Start
    doc.add_heading('2. Quick Start', level=1)

    p = doc.add_paragraph()
    p.add_run('Web Access Instructions').bold = True
    p.add_run('\nLocal URL: http://127.0.0.1:8000/')

    doc.add_heading('Registration Process', level=2)
    steps = [
        'Visit the system and click "Register" button (Local URL: http://127.0.0.1:8000/)',
        'Enter your 11-digit Chinese mobile number',
        'Request verification code (in test mode, it prints to console)',
        'Set password (minimum 8 characters with letters and numbers)',
        'Submit to complete registration',
    ]

    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.add_run(f'Step {i+1}: ').bold = True
        p.add_run(step)

    doc.add_heading('Login & Vote Submission', level=2)

    steps = [
        'Enter your mobile number and password on the login page',
        'Click "Login" to receive authentication Token (valid 24 hours)',
        'Go to "Vote" section and select target category',
        'Search company name, fill voting reason, upload evidence',
        'View leaderboard after submission',
    ]

    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.add_run(f'Step {i+1}: ').bold = True
        p.add_run(step)

    doc.add_page_break()

    # Section 3: Feature Modules
    doc.add_heading('3. Feature Modules', level=1)

    doc.add_heading('API Endpoints Overview', level=2)

    api_list = [
        'POST /api/auth/register/ - Register with phone and verification code',
        'POST /api/auth/login/ - Login to obtain authentication token',
        'POST /api/vote/submit/ - Submit vote for observation category',
        'GET /api/vote/ranking/<category>/ - Query specific category leaderboard',
        'GET /api/vote/companies/search/?q=<keyword> - Search enterprises by name',
        'POST /api/vote/evidence/upload/ - Upload supporting evidence materials',
    ]

    for api in api_list:
        doc.add_paragraph(api, style='List Bullet')

    doc.add_page_break()

    # Section 4: API Documentation Details
    doc.add_heading('4. API Documentation Details', level=1)

    p = doc.add_paragraph()
    p.add_run('Authentication Token Format').bold = True
    p.add_run('\nAuthorization: Token <your_token_key>\n')

    p = doc.add_paragraph()
    p.add_run('Supported File Types for Evidence Upload:').bold = True

    supported_formats = [
        'Images: .jpg, .jpeg, .png',
        'Documents: .pdf, .docx',
        'Max file size: 5MB per file',
    ]
    for fmt in supported_formats:
        doc.add_paragraph(fmt, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Valid Category Values:').bold = True

    categories = [
        'age_discrimination - 35+ Age Discrimination Observation List',
        'gender_discrimination - Gender Discrimination Observation List',
        'pua_despotism - Workplace PUA Despotism Observation List',
        'overwork_exploitation - Overwork Exploitation Observation List',
    ]
    for cat in categories:
        doc.add_paragraph(cat, style='List Bullet')

    # Section 5: FAQ
    doc.add_heading('5. FAQ & Troubleshooting', level=1)

    p = doc.add_paragraph()
    p.add_run('Q1: Daily vote limit reached?').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('A: Same user can cast max 3 votes per day. Solution: Continue using tomorrow after midnight, or register new account.\n')

    p = doc.add_paragraph()
    p.add_run('Q2: Already voted on a company - can I modify or delete?').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('A: Current version does NOT support modifying or deleting votes. Contact administrator for manual handling.\n')

    p = doc.add_paragraph()
    p.add_run('Q3: Evidence upload fails?').bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('A: Check file format (only .jpg/.jpeg/.png/.pdf/.docx supported), file size (<5MB), and avoid duplicate submissions.\n')

    # Appendix A: Test Scripts Summary
    doc.add_heading('Appendix A: Test Script Examples', level=1)

    doc.add_paragraph('Python Requests (End-to-End Test)', style='Heading 3')
    test_snippet_1 = '''login_resp = requests.post(BASE_URL + '/api/auth/login/', json={'username': 'testuser', 'password': 'TestPass123!'})
token = login_resp.json()['token']

vote_resp = s.post(BASE_URL + '/api/vote/submit/', data={
    'category': 'age_discrimination',
    'company_id': companies[0]['id'],
    'reason': 'Evidence for age discrimination'
})'''
    doc.add_paragraph(test_snippet_1)

    doc.add_paragraph('cURL Command Examples', style='Heading 3')
    test_snippet_2 = '''curl -X POST http://127.0.0.1:8000/api/auth/login/ \\
  -H "Content-Type: application/json" \\
  -d '{"username":"testuser","password":"TestPass123!"}'

curl -s http://127.0.0.1:8000/api/vote/companies/search/?q=Huawei'''
    doc.add_paragraph(test_snippet_2)

    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), 'English User Manual.docx')
    doc.save(output_path)
    print(f'[OK] English User Manual Word saved: {output_path}')


def create_en_pdf():
    """Create English User Manual as PDF."""

    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4

    pdf_file = os.path.join(os.path.dirname(__file__), 'English User Manual.pdf')
    doc_build = SimpleDocTemplate(pdf_file, pagesize=A4)

    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle('EN_H1', parent=styles['Heading1'], fontSize=20)
    h2_style = ParagraphStyle('EN_H2', parent=styles['Heading2'], fontSize=16)
    body_style = ParagraphStyle('EN_Body', parent=styles['Normal'], fontSize=9, leading=8)

    elements = []

    # Title
    elements.append(Paragraph('Earth Citizens Supervision Alliance - User Guide (English)', h1_style))
    elements.append(Spacer(1, 20))

    sections = [
        ('System Overview', 'Netizen-driven platform for monitoring employment discrimination through anonymous voting with evidence submission.'),
        ('Observation Categories', '35+ Age Discrimination | Gender Discrimination | Workplace PUA Despotism | Overwork Exploitation'),
        ('Quick Start Guide', '1) Register with phone number + verification code\n2) Login to obtain authentication token\n3) Submit votes and upload evidence materials\n4) View real-time leaderboard rankings'),
        ('Core API Endpoints', '/api/auth/register/ | /api/auth/login/ | /api/vote/submit/ | /api/vote/ranking/ | /api/vote/companies/search/ | /api/vote/evidence/upload/'),
    ]

    for title, content in sections:
        elements.append(Paragraph(title, h2_style))
        elements.append(Spacer(1, 10))
        for line in content.split('\n'):
            if line.strip():
                elements.append(Paragraph(line.strip(), body_style))

    doc_build.build(elements)
    print(f'[OK] English User Manual PDF saved: {pdf_file}')


def main():
    """Main entry point."""
    from reportlab.platypus import SimpleDocTemplate, Spacer
    from reportlab.lib.pagesizes import A4

    print('='*60)
    print('English Document Generator - Earth Citizens Supervision Alliance')
    print('='*60)

    create_en_word_docx()
    create_en_pdf()

    print('\n' + '='*60)
    print('All English documents successfully generated!')
    print('='*60)


if __name__ == '__main__':
    main()
